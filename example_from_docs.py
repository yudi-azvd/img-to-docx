from docx import Document # documentation suggestion
# from docx.document import Document # to make intellisense work
from docx.shared import Inches

document = Document()

document.add_heading('Document YUDI', 0)

par = document.add_paragraph('A plain text having')
par.add_run('bold').bold = True
par.add_run('and some ').bold = True
par.add_run('italic').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph('first item in unordered list', style='List Bullet')
document.add_paragraph('first item in ordered list', style='List Number')

document.add_picture('imgs\IMG_20180720_143234.jpg', width=Inches(1.25))

records = (
  (3, '101', 'Spam'),
  (7, '422', 'Eggs'),
  (4, '631', 'Spam spam eggs and spam'),
)

table = document.add_table(rows=1, cols=3)
header_cells = table.rows[0].cells
header_cells[0].text = 'Qty'
header_cells[1].text = 'Id'
header_cells[2].text = 'Desc'

for qty, id, desc in records:
  row_cells = table.add_row().cells
  row_cells[0].text = str(qty)
  row_cells[1].text = id
  row_cells[2].text = desc

document.add_page_break()

document.sections[0]

document.save('demo.docx')
