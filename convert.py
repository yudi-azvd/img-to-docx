# -*- coding: utf-8 -*-

import os
import glob

from docx import Document # documentation suggestion
# from docx.document import Document # to make intellisense work
from docx.oxml.xmlchemy import OxmlElement
from docx.oxml.shared import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm


def add_content_to_doc(paths_to_imgs):
  document = Document()

  for path_to_img in paths_to_imgs:
    add_header(document)
    add_pic(document, path_to_img)
    config_margins(document)
    config_borders(document)

  document.save('Atividades de Matemática.docx')

  return


def add_header(document):
  document.add_paragraph('\tEscola Classe 303 de São Sebastião')
  document.add_paragraph('\tNome completo: ___________________________________________________________________    Professora: _________________________')
  document.add_paragraph('\tSérie: _____________________________ \t                                                Data: ______/______/______.')
  par = document.add_paragraph('')
  par.alignment = WD_ALIGN_PARAGRAPH.CENTER
  par.add_run('Matemática').bold = True
  return


def add_pic(document, path_to_img):
  document.add_picture(path_to_img, width=Cm(16.5))
  par_with_pic = document.paragraphs[-1]
  par_with_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
  return


# Assuming there is only one section in document
def config_margins(document):
  section = document.sections[0]
  margin = Cm(0.5)
  section.left_margin = margin
  section.right_margin = margin
  section.top_margin = margin
  section.bottom_margin = Cm(0)

  return

# Huge thanks to this post
# https://stackoverflow.com/questions/55783519/how-to-add-a-page-border-with-python-docx
# Documentation: http://officeopenxml.com/WPsectionBorders.php
def config_borders(document: Document):
  section_properties = document.sections[0]._sectPr
  page_borders = OxmlElement('w:pgBorders')
  page_borders.set(qn('w:offsetFrom'), 'page')

  for border_name in ('top', 'bottom'):
    border = OxmlElement(f'w:{border_name}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), '8')
    border.set(qn('w:space'), '10')
    border.set(qn('w:color'), 'auto')
    page_borders.append(border)

  for border_name in ('right', 'left'):
    border = OxmlElement(f'w:{border_name}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), '8')
    border.set(qn('w:space'), '15')
    border.set(qn('w:color'), 'auto')
    page_borders.append(border)

  section_properties.append(page_borders)
  return


def run():
  paths_to_imgs = glob.glob(os.path.join('imgs', '*.jpg'))
  add_content_to_doc(paths_to_imgs)
  return

run()
