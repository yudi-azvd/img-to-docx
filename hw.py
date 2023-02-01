#!/usr/bin/env python

from docx import Document
from docx.oxml.xmlchemy import OxmlElement
from docx.oxml.shared import qn
from docx.shared import Cm


def create_text_box(doc: Document, name: str):
    wp = OxmlElement('w:p')

    wr = OxmlElement('w:r')
    wrpr = OxmlElement('w:rPr')
    wsz = OxmlElement('w:sz')
    wsz.set(qn('w:val'), '24')
    wszCs = OxmlElement('w:szCs')
    wszCs.set(qn('w:val'), '24')
    wrpr.append(wsz)
    wrpr.append(wszCs)
    text = OxmlElement('w:t')
    text.text = f'Olaaaaaa {name}'
    wr.append(wrpr)
    wr.append(text)
    wp.append(wr)
    return wp


def create_page(doc: Document, name: str):
    section = doc.sections[0]
    section.page_width = Cm(29.70)
    section.page_height = Cm(21.00)
    section_properties = doc.sections[0]._sectPr
    page_borders = OxmlElement('w:pgBorders')
    page_borders.set(qn('w:offsetFrom'), 'page')

    for border_name in ('top', 'bottom'):
        border = OxmlElement('w:'+border_name)
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')
        border.set(qn('w:space'), '10')
        border.set(qn('w:color'), 'auto')
        page_borders.append(border)

    for border_name in ('right', 'left'):
        border = OxmlElement('w:'+border_name)
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')
        border.set(qn('w:space'), '15')
        border.set(qn('w:color'), 'auto')
        page_borders.append(border)

    section_properties.append(page_borders)

    margin = Cm(1)
    section.left_margin = margin
    section.right_margin = margin
    section.top_margin = margin
    section.bottom_margin = margin

    doc.add_picture('img/numeros.png', width=Cm(27.0))
    doc.add_paragraph(name)
    section_properties.append(create_text_box(doc, name))
    doc.add_picture('img/bocas.png', width=Cm(27.0))

    return


def main():
    document = Document()

    students_names = [
        'yudi'
    ]

    for name in students_names:
        create_page(document, name)

    document.save('atividade.docx')
    return


if __name__ == '__main__':
    main()
